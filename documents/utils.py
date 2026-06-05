import logging
import re

import pypdf
from docx import Document as DocxDocument
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


def clean_extracted_text(text):
    """Очищает извлечённый текст для улучшения поиска"""
    if not text:
        return ""

    # 1. Заменяем NUL символы на пробелы (важно для склеивания слов)
    text = text.replace("\x00", " ")

    # 2. Объединяем разорванные слова (было: "ра зрыв" → "разрыв")
    text = re.sub(r"(\w)\s+(\w)", r"\1\2", text)

    # 3. Убираем лишние пробелы (более одного подряд)
    text = re.sub(r"[ \t]+", " ", text)

    # 4. Убираем пробелы вокруг пунктуации
    text = re.sub(r"\s+([.,!?;:])", r"\1", text)

    # 5. Нормализуем переносы строк (3+ переносов → 2)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # 6. Убираем пробелы в начале/конце строк
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)

    return text.strip()


def extract_text_from_file(file_path, file_type):
    """Извлекает текст из файла в зависимости от его типа. PDF, DOCX, XLSX, TXT"""
    text = ""

    try:
        if file_type == "pdf":
            # Извлечение текста из PDF
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

        elif file_type == "docx":
            # Извлечение текста из DOCX
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Извлекаем текст из таблиц, если они есть
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

        elif file_type == "xlsx":
            # Извлечение текста из XLSX
            wb = load_workbook(file_path, data_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    # Преобразуем все ячейки в строки и объединяем
                    text += " ".join([str(cell) for cell in row if cell]) + "\n"

        elif file_type == "txt":
            # Прямое чтение текстового файла
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}", exc_info=True)
        return ""

    # Применяем очистку текста
    text = clean_extracted_text(text)

    return text
